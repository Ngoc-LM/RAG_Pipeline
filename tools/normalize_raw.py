"""Trích và chuẩn hoá văn bản QPPL từ PDF/DOC/DOCX thành text nạp được vào corpus.

Toàn bộ logic của bước chuẩn bị corpus nằm ở đây chứ không nằm trong notebook:
mã trong cell notebook không test được, không diff được và không tái hiện được.
`notebooks/prepare_corpus.py` chỉ upload file, gọi các hàm dưới đây và in báo cáo.

Đầu ra của `normalize()` phải khớp giả định của `src/ingest.py`: ranh giới cấu
trúc nằm ở đầu dòng, số khoản tăng liên tiếp, không còn quốc hiệu hay phụ lục.
"""

from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
import unicodedata
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING, Iterable, Mapping, Sequence

from src.chunk import estimate_tokens
from src.ingest import REQUIRED_KEYS, VALID_STATUS

if TYPE_CHECKING:
    from src.chunk import Chunk as ChunkObj
    from src.ingest import Document as IngestDocument

CHAPTER_RE = re.compile(r"^Chương\s+[IVXLCDM]+")
SECTION_RE = re.compile(r"^Mục\s+\d+")
ARTICLE_RE = re.compile(r"^Điều\s+(\d+)\.")
CLAUSE_RE = re.compile(r"^\d+\.\s")
POINT_RE = re.compile(r"^[a-zđ]\)\s")
STRUCTURAL = (CHAPTER_RE, SECTION_RE, ARTICLE_RE, CLAUSE_RE, POINT_RE)
HEADINGS = (CHAPTER_RE, SECTION_RE, ARTICLE_RE)
"""Dòng tiêu đề — khác khoản và điểm ở chỗ chúng không bao giờ được nuốt dòng sau."""

ARTICLE_TITLE_RE = re.compile(r"^Điều\s+(\d+)\.\s*(.*)$")
PAGE_NUMBER_RE = re.compile(r"^\d+$")
START_RE = re.compile(r"^(Chương\s+[IVXLCDM]+|Điều\s+1\.)")
TAIL_RE = re.compile(r"^(Nơi nhận|PHỤ LỤC|Phụ lục|Biểu mẫu)\b|^_{5,}")
"""Mốc kết thúc phần quy phạm.

Dãy gạch dưới là dòng kẻ ngăn trước phần công bố ("Luật này được Quốc hội ...
thông qua ngày ..."). Bước nối dòng chạy trước bước cắt nên câu công bố đã dính
vào dòng kẻ, cắt theo dòng kẻ là bỏ được cả hai.
"""
SENTENCE_END = ".;:!?"

REPEATED_LINE_THRESHOLD = 3
"""Dòng giống hệt nhau xuất hiện nhiều hơn ngần này lần là header/footer chạy trang."""

MIN_PDF_CHARS = 500
MIN_DIACRITIC_RATIO = 0.01
LONG_ARTICLE_TOKENS = 2000

MAGIC = {
    b"%PDF": "pdf",
    b"\xd0\xcf\x11\xe0": "doc",
    b"PK\x03\x04": "docx",
}

FRONTMATTER_ORDER = (
    "doc_id",
    "title",
    "doc_type",
    "issued_date",
    "effective_from",
    "effective_to",
    "status",
)
_BARE_VALUE_RE = re.compile(r"^[A-Za-z0-9_]+$")


class ExtractionError(RuntimeError):
    """Không trích được text dùng được từ file nguồn."""


def _is_structural(line: str) -> bool:
    return any(pattern.match(line) for pattern in STRUCTURAL)


# --- trích text -----------------------------------------------------------
def detect_format(path: Path) -> str:
    """Nhận dạng theo magic bytes trước, đuôi file chỉ là phương án dự phòng.

    File .doc nhị phân cũ (OLE2) hay bị đặt tên .docx và ngược lại; python-docx
    chỉ đọc được ZIP nên đoán sai định dạng sẽ nổ ở tận bước đọc với thông báo
    khó hiểu.
    """
    header = Path(path).open("rb").read(8)
    for magic, name in MAGIC.items():
        if header.startswith(magic):
            return name
    suffix = Path(path).suffix.lower().lstrip(".")
    return suffix if suffix in ("pdf", "doc", "docx") else "unknown"


def diacritic_ratio(text: str) -> float:
    """Tỉ lệ ký tự mang dấu tiếng Việt, dùng để phát hiện PDF scan hoặc lỗi font."""
    if not text:
        return 0.0
    marked = sum(
        1
        for character in text
        if character in "đĐ"
        or any(unicodedata.combining(c) for c in unicodedata.normalize("NFD", character))
    )
    return marked / len(text)


def check_extraction_quality(text: str, source: str) -> None:
    """Chặn corpus rỗng hoặc mất dấu ngay tại chỗ thay vì để nó trôi xuống dưới."""
    if len(text) < MIN_PDF_CHARS:
        raise ExtractionError(
            f"{source}: chỉ trích được {len(text)} ký tự (< {MIN_PDF_CHARS}). "
            "PDF nhiều khả năng là bản scan, cần bản .docx thay thế; "
            "OCR nằm ngoài phạm vi."
        )
    ratio = diacritic_ratio(text)
    if ratio < MIN_DIACRITIC_RATIO:
        raise ExtractionError(
            f"{source}: chỉ {ratio:.2%} ký tự có dấu tiếng Việt "
            f"(< {MIN_DIACRITIC_RATIO:.0%}). PDF nhiều khả năng là bản scan, "
            "cần bản .docx thay thế; OCR nằm ngoài phạm vi."
        )


def _run(command: Sequence[str], what: str) -> None:
    if shutil.which(command[0]) is None:
        raise ExtractionError(f"{what}: không tìm thấy `{command[0]}` trên PATH")
    result = subprocess.run(command, capture_output=True, text=True, timeout=300)
    if result.returncode != 0:
        raise ExtractionError(f"{what}: `{command[0]}` lỗi {result.returncode}\n{result.stderr}")


def extract_pdf(path: Path) -> str:
    """pdftotext -layout giữ được thứ tự đọc của văn bản một cột."""
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / "out.txt"
        _run(
            ["pdftotext", "-layout", "-nopgbrk", "-enc", "UTF-8", str(path), str(out)],
            f"trích PDF {path.name}",
        )
        return out.read_text(encoding="utf-8", errors="replace")


def convert_doc_to_docx(path: Path, outdir: Path) -> Path:
    """Chuyển .doc nhị phân cũ sang .docx bằng LibreOffice headless.

    Không dùng antiword: nó hỏng dấu tiếng Việt. python-docx chỉ đọc được ZIP
    nên .doc OLE2 bắt buộc phải qua bước chuyển đổi này.
    """
    outdir.mkdir(parents=True, exist_ok=True)
    _run(
        ["soffice", "--headless", "--convert-to", "docx", "--outdir", str(outdir), str(path)],
        f"chuyển {path.name} sang docx",
    )
    converted = outdir / f"{path.stem}.docx"
    if not converted.is_file():
        raise ExtractionError(f"{path.name}: soffice chạy xong nhưng không thấy {converted.name}")
    return converted


def extract_docx(path: Path) -> str:
    """Nối các paragraph bằng "\\n", giữ paragraph rỗng làm ranh giới đoạn.

    Bỏ qua nội dung trong bảng: `Document.paragraphs` chỉ trả về paragraph ở thân
    tài liệu, không đi vào ô bảng. Bảng biểu trong văn bản QPPL không parse được
    theo cấu trúc Điều/Khoản và không sinh câu hỏi tốt.
    """
    from docx import Document as DocxDocument

    document = DocxDocument(str(path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def extract_text(path: Path) -> str:
    """Trích text thô từ .pdf / .doc / .docx, đã kiểm chất lượng tối thiểu."""
    path = Path(path)
    kind = detect_format(path)

    if kind == "pdf":
        text = extract_pdf(path)
    elif kind == "docx":
        text = extract_docx(path)
    elif kind == "doc":
        with tempfile.TemporaryDirectory() as tmp:
            text = extract_docx(convert_doc_to_docx(path, Path(tmp)))
    else:
        raise ExtractionError(
            f"{path.name}: không nhận dạng được định dạng (magic bytes lẫn đuôi file)"
        )

    check_extraction_quality(text, path.name)
    return text


# --- chuẩn hoá ------------------------------------------------------------
@dataclass
class NormalizeReport:
    n_lines_joined: int = 0
    n_lines_dropped: int = 0
    n_chapters: int = 0
    n_articles: int = 0
    n_clauses: int = 0
    warnings: list[str] = field(default_factory=list)

    def summary(self) -> str:
        head = (
            f"nối {self.n_lines_joined} dòng, bỏ {self.n_lines_dropped} dòng | "
            f"{self.n_chapters} chương, {self.n_articles} điều, {self.n_clauses} khoản"
        )
        if not self.warnings:
            return head
        return head + "\n" + "\n".join(f"  ⚠ {w}" for w in self.warnings)


_CHAR_FIXES = {
    "﻿": "",
    " ": " ",
    "­": "",
    "‘": "'",
    "’": "'",
    "‚": "'",
    "‛": "'",
    "“": '"',
    "”": '"',
    "„": '"',
    "‟": '"',
    "–": "-",
    "—": "-",
    "‒": "-",
    "―": "-",
}


def _clean_characters(text: str) -> str:
    text = unicodedata.normalize("NFC", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    for source, target in _CHAR_FIXES.items():
        text = text.replace(source, target)
    return text


def _drop_noise_lines(lines: list[str], report: NormalizeReport) -> list[str]:
    """Bỏ số trang và header/footer chạy trang."""
    counts = Counter(line.strip() for line in lines if line.strip())
    repeated = {
        line for line, count in counts.items() if count > REPEATED_LINE_THRESHOLD
    }
    if repeated:
        report.warnings.append(
            f"bỏ {len(repeated)} dòng lặp (>{REPEATED_LINE_THRESHOLD} lần): "
            + ", ".join(sorted(repeated)[:3])[:120]
        )

    kept: list[str] = []
    for line in lines:
        stripped = line.strip()
        if PAGE_NUMBER_RE.match(stripped) or (stripped and stripped in repeated):
            report.n_lines_dropped += 1
            continue
        kept.append(line)
    return kept


def _join_wrapped(lines: list[str], report: NormalizeReport) -> list[str]:
    """Nối dòng bị PDF ngắt cứng giữa câu.

    Bước quan trọng nhất: nếu bỏ qua, mỗi dòng vật lý thành một "đoạn" và
    `src/ingest.py` sẽ thấy khoản bị vỡ vụn. Chỉ nối khi dòng hiện tại không mở
    một đơn vị cấu trúc mới và dòng trước không kết thúc bằng dấu câu.

    Thêm một chốt không có trong quy tắc gốc: dòng trước là tiêu đề Chương/Mục/
    Điều thì không nối. Tiêu đề gần như không bao giờ kết thúc bằng dấu câu, nên
    nếu thiếu chốt này thì "Điều 1. Phạm vi điều chỉnh" sẽ nuốt câu đầu của thân
    điều (ingest lấy phần sau dấu chấm làm article_title, và Điều mất sạch nội
    dung), còn "Chương I" nuốt "QUY ĐỊNH CHUNG" khiến CHAPTER_RE của ingest —
    vốn neo `$` ngay sau số La Mã — không khớp nữa và cả chương biến mất.
    Đánh đổi: tiêu đề Điều dài bị PDF ngắt làm hai dòng sẽ không được nối lại.
    """
    joined: list[str] = []
    for line in lines:
        stripped = line.strip()
        previous = joined[-1].strip() if joined else ""
        if (
            stripped
            and previous
            and not _is_structural(stripped)
            and not any(pattern.match(previous) for pattern in HEADINGS)
            and previous[-1] not in SENTENCE_END
        ):
            joined[-1] = f"{joined[-1].rstrip()} {stripped}"
            report.n_lines_joined += 1
            continue
        joined.append(line)
    return joined


def _tidy_whitespace(lines: list[str]) -> list[str]:
    """Gộp space thừa, bỏ space cuối dòng, gộp mọi chuỗi dòng trống thành đúng một."""
    squeezed = [re.sub(r"[ \t]+", " ", line).rstrip() for line in lines]
    tidied: list[str] = []
    for line in squeezed:
        if not line and tidied and not tidied[-1]:
            continue
        tidied.append(line)
    return tidied


def _cut_scope(lines: list[str], report: NormalizeReport) -> list[str]:
    """Bỏ quốc hiệu, tiêu ngữ, phần "Căn cứ ..." ở đầu và phụ lục ở cuối."""
    start = next((i for i, line in enumerate(lines) if START_RE.match(line.strip())), None)
    if start is None:
        report.warnings.append(
            "không tìm thấy dòng ^Chương hoặc ^Điều 1. — giữ nguyên phần đầu, cần soi tay"
        )
        start = 0
    else:
        report.n_lines_dropped += start

    end = next(
        (i for i, line in enumerate(lines[start:], start) if TAIL_RE.match(line.strip())),
        None,
    )
    if end is not None:
        report.n_lines_dropped += len(lines) - end
        report.warnings.append(f"cắt từ dòng {end} trở đi: {lines[end].strip()[:60]!r}")
        return lines[start:end]
    return lines[start:]


def _count_structure(lines: Sequence[str], report: NormalizeReport) -> None:
    numbers: list[int] = []
    for line in lines:
        stripped = line.strip()
        if CHAPTER_RE.match(stripped):
            report.n_chapters += 1
        match = ARTICLE_RE.match(stripped)
        if match:
            report.n_articles += 1
            numbers.append(int(match.group(1)))
        if CLAUSE_RE.match(stripped):
            report.n_clauses += 1

    if not numbers:
        report.warnings.append("không nhận được Điều nào — nhiều khả năng dòng 'Điều N.' bị wrap")
        return
    gaps = [
        (a, b) for a, b in zip(numbers, numbers[1:]) if b != a + 1
    ]
    if gaps:
        report.warnings.append(
            "số Điều không liên tiếp: " + ", ".join(f"{a}->{b}" for a, b in gaps[:5])
        )


def normalize(text: str) -> tuple[str, NormalizeReport]:
    """Chuẩn hoá text thô thành dạng `src/ingest.py` parse được."""
    report = NormalizeReport()
    lines = _clean_characters(text).split("\n")
    lines = _drop_noise_lines(lines, report)
    lines = _join_wrapped(lines, report)
    lines = _tidy_whitespace(lines)
    lines = _cut_scope(lines, report)
    lines = _tidy_whitespace(lines)
    _count_structure(lines, report)
    return "\n".join(lines).strip() + "\n", report


# --- soi cấu trúc ---------------------------------------------------------
@dataclass(frozen=True)
class OutlineArticle:
    chapter: str | None
    article_no: int
    title: str
    n_clauses: int
    n_tokens: int
    warnings: tuple[str, ...]


def outline_tree(text: str) -> list[OutlineArticle]:
    """Cây Chương/Điều kèm cảnh báo, để soi TRƯỚC khi quyết định cắt phạm vi."""
    lines = text.split("\n")
    entries: list[OutlineArticle] = []
    chapter: str | None = None
    current: dict[str, object] | None = None

    def close(pending: dict[str, object] | None) -> None:
        if pending is None:
            return
        clause_numbers: list[int] = pending["clause_numbers"]  # type: ignore[assignment]
        body = "\n".join(pending["body"])  # type: ignore[arg-type]
        tokens = estimate_tokens(body)
        warnings: list[str] = []
        if not clause_numbers:
            warnings.append("0 khoản")
        elif clause_numbers != list(range(1, len(clause_numbers) + 1)):
            warnings.append(f"khoản nhảy cóc: {clause_numbers}")
        if tokens > LONG_ARTICLE_TOKENS:
            warnings.append(
                f"dài {tokens} token — nhiều khả năng dòng 'Điều N.' kế tiếp bị wrap"
            )
        entries.append(
            OutlineArticle(
                chapter=pending["chapter"],  # type: ignore[arg-type]
                article_no=pending["article_no"],  # type: ignore[arg-type]
                title=pending["title"],  # type: ignore[arg-type]
                n_clauses=len(clause_numbers),
                n_tokens=tokens,
                warnings=tuple(warnings),
            )
        )

    for index, line in enumerate(lines):
        stripped = line.strip()
        if CHAPTER_RE.match(stripped):
            close(current)
            current = None
            following = next(
                (
                    lines[j].strip()
                    for j in range(index + 1, len(lines))
                    if lines[j].strip()
                ),
                "",
            )
            chapter = (
                f"{stripped} — {following}"
                if following and not _is_structural(following)
                else stripped
            )
            continue

        match = ARTICLE_TITLE_RE.match(stripped)
        if match:
            close(current)
            current = {
                "chapter": chapter,
                "article_no": int(match.group(1)),
                "title": match.group(2).strip(),
                "clause_numbers": [],
                "body": [],
            }
            continue

        if current is not None:
            current["body"].append(line)  # type: ignore[union-attr]
            clause = CLAUSE_RE.match(stripped)
            if clause:
                current["clause_numbers"].append(int(stripped.split(".", 1)[0]))  # type: ignore[union-attr]

    close(current)
    return entries


def outline(text: str) -> str:
    """Cây cấu trúc dạng chuỗi in ra được."""
    lines: list[str] = []
    chapter: str | None = "\x00"
    for entry in outline_tree(text):
        if entry.chapter != chapter:
            chapter = entry.chapter
            lines.append(chapter or "(ngoài chương)")
        label = f"  Điều {entry.article_no}. {entry.title}"
        stats = f"({entry.n_clauses} khoản, {entry.n_tokens} token)"
        lines.append(f"{label:<52} {stats}")
        for warning in entry.warnings:
            lines.append(f"{'':<52} ⚠ {warning}")
    return "\n".join(lines) if lines else "(không nhận được Điều nào)"


# --- cắt phạm vi ----------------------------------------------------------
def _wanted_numbers(keep: Mapping[str, Iterable[int]] | Iterable[int]) -> set[int]:
    if isinstance(keep, Mapping):
        return {int(n) for n in keep.get("articles", ())}
    return {int(n) for n in keep}


def select_articles(
    text: str, keep: Mapping[str, Iterable[int]] | Iterable[int]
) -> str:
    """Giữ lại các Điều được chọn cùng dòng tiêu đề Chương chứa chúng.

    Không đánh số lại: số Điều gốc là thứ người đọc dùng để tra ngược văn bản,
    đánh số lại sẽ làm mọi trích dẫn sai.
    """
    wanted = _wanted_numbers(keep)
    lines = text.split("\n")
    out: list[str] = []
    pending: dict[str, list[str] | None] = {"chapter": None, "section": None}

    def capture(start: int) -> tuple[list[str], int]:
        block = [lines[start]]
        cursor = start + 1
        while (
            cursor < len(lines)
            and lines[cursor].strip()
            and not _is_structural(lines[cursor].strip())
        ):
            block.append(lines[cursor])
            cursor += 1
        return block, cursor

    index = 0
    while index < len(lines):
        stripped = lines[index].strip()

        if CHAPTER_RE.match(stripped):
            pending["chapter"], index = capture(index)
            pending["section"] = None
            continue

        if SECTION_RE.match(stripped):
            pending["section"], index = capture(index)
            continue

        match = ARTICLE_RE.match(stripped)
        if match:
            cursor = index + 1
            while cursor < len(lines) and not (
                CHAPTER_RE.match(lines[cursor].strip())
                or SECTION_RE.match(lines[cursor].strip())
                or ARTICLE_RE.match(lines[cursor].strip())
            ):
                cursor += 1
            if int(match.group(1)) in wanted:
                # Tiêu đề Chương/Mục chỉ được giữ khi thực sự còn Điều nằm dưới
                # nó, nếu không sẽ thành tiêu đề mồ côi và tạo lỗ thủng coverage.
                for level in ("chapter", "section"):
                    block = pending[level]
                    if block is not None:
                        if out:
                            out.append("")
                        out.extend(block)
                        out.append("")
                        pending[level] = None
                out.extend(lines[index:cursor])
            index = cursor
            continue

        index += 1

    return "\n".join(out).strip() + "\n"


def word_count(text: str) -> int:
    return len(text.split())


# --- ghi corpus -----------------------------------------------------------
def render_frontmatter(meta: Mapping[str, str | None]) -> str:
    """Sinh khối frontmatter đúng dạng `src/ingest.py` đọc được.

    Kiểm ngay tại đây bằng REQUIRED_KEYS và VALID_STATUS nhập từ src.ingest, để
    hai đầu không trôi khỏi nhau mà phải chạy cả pipeline mới phát hiện.
    """
    missing = [key for key in REQUIRED_KEYS if not meta.get(key)]
    if missing:
        raise ValueError(f"frontmatter thiếu key bắt buộc {missing}")
    if meta.get("status") not in VALID_STATUS:
        raise ValueError(f"status={meta.get('status')!r} phải thuộc {VALID_STATUS}")

    rows = ["---"]
    for key in FRONTMATTER_ORDER:
        value = meta.get(key)
        if value is None:
            rendered = "null"
        elif _BARE_VALUE_RE.match(str(value)):
            rendered = str(value)
        else:
            rendered = '"{}"'.format(str(value).replace('"', "'"))
        rows.append(f"{key}: {rendered}")
    rows.append("---")
    return "\n".join(rows) + "\n"


def corpus_audit(
    documents: Sequence["IngestDocument"],
    chunks: Sequence["ChunkObj"],
    keep: Mapping[str, Mapping[str, Iterable[int]]] | None = None,
) -> str:
    """Bảng kiểm khép vòng cho corpus vừa ghi ra.

    Đứng ở tools/ chứ không ở notebook để cell chỉ còn hai dòng gọi hàm, và để
    chính bảng kiểm này cũng test được.
    """
    from statistics import median

    from src.intervals import interval_union

    rows = [
        f"{'doc_id':<22} {'điều':>5} {'chunk':>6} {'median tok':>11} {'phủ body':>9}",
        "-" * 58,
    ]
    problems: list[str] = []

    for document in documents:
        own = [c for c in chunks if c.doc_id == document.meta.doc_id]
        covered = sum(
            end - start
            for start, end in interval_union([(c.char_start, c.char_end) for c in own])
        )
        ratio = covered / len(document.body) if document.body else 0.0
        tokens = [c.n_tokens for c in own] or [0]
        rows.append(
            f"{document.meta.doc_id:<22} {len(document.articles):>5} {len(own):>6} "
            f"{median(tokens):>11.0f} {ratio:>8.1%}"
        )

        if keep is None:
            continue
        wanted = _wanted_numbers(keep.get(document.meta.doc_id, {}))
        parsed = {a.article_no for a in document.articles}
        if wanted and wanted != parsed:
            missing = sorted(wanted - parsed)
            extra = sorted(parsed - wanted)
            problems.append(
                f"LỆCH {document.meta.doc_id}: KEEP có {len(wanted)} Điều, "
                f"parser nhận {len(parsed)}"
                + (f" | thiếu {missing}" if missing else "")
                + (f" | thừa {extra}" if extra else "")
            )

    if problems:
        rows.append("")
        rows.extend(problems)
    return "\n".join(rows)


def write_corpus_file(directory: Path, meta: Mapping[str, str | None], body: str) -> Path:
    """Ghi data/corpus/<doc_id>.txt với LF và UTF-8."""
    directory = Path(directory)
    directory.mkdir(parents=True, exist_ok=True)
    path = directory / f"{meta['doc_id']}.txt"
    content = f"{render_frontmatter(meta)}\n{body.strip()}\n"
    path.write_text(content, encoding="utf-8", newline="\n")
    return path
